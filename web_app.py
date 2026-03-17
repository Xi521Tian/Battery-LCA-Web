import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
import io
import datetime
import re

st.set_page_config(page_title="电池LCA在线测算", page_icon="🔋", layout="wide")
st.title("🔋 动力电池全生命周期 (LCA) 在线核算与报告系统")
st.markdown("---")

# ==========================================
# 🌟 第一步：配置报告基础信息
# ==========================================
st.sidebar.header("📝 填报导航")
st.sidebar.info("请依次完成基础信息配置与数据录入，系统将自动生成完整的带封面、目录的正式报告。")

st.header("第一步：配置报告基础信息")

with st.expander("📂 Part 1: 概述 (带*号为必填项)", expanded=True):
    company_intro = st.text_area("* 1.1 企业简介", placeholder="请输入企业基本情况（必填）...")
    product_intro = st.text_area("* 1.2 产品介绍", placeholder="请输入产品型号、规格（必填）...")
    production_process = st.text_area("* 1.3 产品生产工艺", placeholder="简述生产制造流程（必填）...")

    def_p1_4 = "基于LCA的评价方法，本产品碳足迹评价报告主要依据了以下标准：国际标准化组织（ISO）发布的《ISO 14067：2018温室气体—产品碳足迹—量化要求与指南》；英国标准协会（BSI）发布的《PAS 2050：2011商品和服务生命周期温室气体排放评价规范》；世界资源研究所（WRI）发布的《温室气体核算体系：产品寿命周期核算与报告标准》。\n依据产品所属行业标准对产品种类规则（Product Category Rules, PCR）的要求，根据体验账号的产品管理规定来定义产品的功能单位、边界、分配等计算原则。"
    evaluation_basis = st.text_area("1.4 评价依据和要求", value=def_p1_4)

def_p2_1 = "本报告旨在通过揭示体验账号生产的汽车动力电池全生命周期的产品碳足迹，为体验账号持续开展节能减排工作提供数据参考；为体验账号向价值链相关方开展碳信息披露提供重要内容。"
def_p2_2 = "本次研究的对象为“汽车动力电池，100kwh”，为了方便输入/输出的量化，保证碳足迹分析结果的可比性，功能单位定义为“1个汽车动力电池”。"
def_p2_3 = "本次评价汽车动力电池的产品碳足迹，其生产周期为2026年01月01日至2026年12月31日。"
def_p2_4 = "本次产品碳足迹评价的系统边界为全生命周期-从资源开采到产品废弃，碳足迹核算包括原材料获取阶段、生产制造阶段、分销和储存阶段、产品使用阶段、废弃处置阶段的温室气体排放或清除。\n本次评估不涉及土地利用变化所导致的温室气体排放或清除，以及产品生物固碳。"
def_p2_5 = "本产品碳足迹评价报告核算二氧化碳（CO2）、甲烷（CH4）、氧化亚氮（N2O）、氢氟碳化物（HFCs）、全氟碳化物（PFCs）、六氟化硫（SF6）和三氟化氮（NF3）向大气中的排放或清除，采用IPCC第六次评估报告100-year的GWP值，将不同温室气体折算为二氧化碳当量（CO2e）。"
def_p2_6 = "本次产品碳足迹评价包括系统边界内所有对产品生命周期温室气体排放具有实质性贡献的排放源。本报告采取的数据取舍原则，以原材料投入占产品重量的比例、排放贡献占总排放的比例为依据，具体原则如下：\n（1）普通物料重量小于产品重量1%时，含有稀贵金属的物料小于产品重量0.1%时，可以忽略，但总共忽略的物料不超过产品重量的5%。\n（2）对于一些较难获取的数据，若其对于排放的影响很小，可对其进行忽略。\n（3）仅考虑生产过程的排放，生活过程不予考虑。\n（4）优先采用供应量最大的供应商数据。\n（5）运输距离无法获取时，采用平均运输距离。"

with st.expander("📂 Part 2: 目的和范围定义 (不填则自动使用默认标准文本)", expanded=False):
    purpose = st.text_area("2.1 研究目的", placeholder=def_p2_1)
    functional_unit = st.text_area("2.2 功能单位", placeholder=def_p2_2)
    time_scope = st.text_area("2.3 时间范围", placeholder=def_p2_3)
    system_boundary = st.text_area("2.4 系统边界", placeholder=def_p2_4)
    ghg_types = st.text_area("2.5 温室气体种类", placeholder=def_p2_5)
    cutoff_rules = st.text_area("2.6 取舍原则", placeholder=def_p2_6)

st.markdown("---")
st.header("第二步：录入生命周期测算数据")

# ==========================================
# 🌟 第二步：因子库与预设 UI 结构
# ==========================================
FACTOR_DB = {
    "天然气 (m3)": 2.0667, "厂务电力 (kWh)": 0.6205, "水 (m3)": 0.344, "废水 (t)": 0.118,
    "正极材料-磷酸铁锂 (kg)": 25.0, "负极材料-石墨 (kg)": 5.5, "电解液 (kg)": 19.6, "隔膜 (kg)": 3.24,
    "铝箔 (kg)": 2.39, "铜箔 (kg)": 12.4, "电池壳体-铝合金 (kg)": 28.38, "电池管理系统BMS (kg)": 28.38,
    "冷却系统-水冷板 (kg)": 11.5, "连接件-铜排 (kg)": 12.4, "电解质锂盐 (kg)": 0.01, "导电剂 (kg)": 3.9,
    "粘结剂 (kg)": 0.002, "电池结构胶 (kg)": 28.38, "绝缘材料 (kg)": 1.85,
    "木质托盘 (kg)": 1.28, "塑料薄膜 (kg)": 3.17, "纸质护角 (个)": 1.13,
    "光伏发电电力 (kWh)": 0.65, "电芯涂布烘干耗电 (kWh)": 0.6205, "辊压工序耗电 (kWh)": 0.6205,
    "电芯注液耗电 (kWh)": 0.6205, "化成分容耗电 (kWh)": 0.6205, "模组焊接耗电 (kWh)": 0.6205,
    "Pack装配耗电 (kWh)": 0.6205, "有机废气 (m3)": 0.056,
    "仓储温控耗电 (kWh)": 0.6205, "车辆行驶充电耗电 (kWh)": 0.6205, "电池热管理系统耗电 (kWh)": 0.6205,
    "回收清洗废水 (t)": 0.858, "动力电池回收拆解耗电 (kWh)": 0.6205,
    "废弃正极材料 (kg)": 1.82, "废弃负极材料 (kg)": 1.82, "废弃电解液 (kg)": 19.6, "废弃隔膜 (kg)": 0.39,
    "废弃铝箔 (kg)": 2.39, "废弃铜箔 (kg)": 12.4, "废弃电池结构胶 (kg)": 28.38, "废弃绝缘材料 (kg)": 1.82,
    "废弃BMS (kg)": 28.38, "废弃铝合金壳体 (kg)": 28.38, "废弃水冷板 (kg)": 0.167, "废弃铜排 (kg)": 0.14
}
default_trans_factor = 0.078

UI_STRUCTURE = {
    "3.5.1 原材料获取阶段": {
        "主辅材与包装": ["正极材料-磷酸铁锂 (kg)", "负极材料-石墨 (kg)", "电解液 (kg)", "隔膜 (kg)", "铝箔 (kg)",
                         "铜箔 (kg)", "电解质锂盐 (kg)", "导电剂 (kg)", "粘结剂 (kg)", "电池壳体-铝合金 (kg)",
                         "电池管理系统BMS (kg)", "冷却系统-水冷板 (kg)", "连接件-铜排 (kg)", "电池结构胶 (kg)",
                         "绝缘材料 (kg)", "木质托盘 (kg)", "塑料薄膜 (kg)", "纸质护角 (个)"],
        "运输": ["正极材料运输 (tkm)", "负极材料运输 (tkm)", "电解液运输 (tkm)", "隔膜运输 (tkm)", "铝箔运输 (tkm)",
                 "铜箔运输 (tkm)", "辅料运输 (tkm)"]
    },
    "3.5.2 生产制造阶段": {
        "物料与能耗": ["厂务电力 (kWh)", "天然气 (m3)", "水 (m3)", "废水 (t)", "有机废气 (m3)", "光伏发电电力 (kWh)",
                       "电芯涂布烘干耗电 (kWh)", "辊压工序耗电 (kWh)", "电芯注液耗电 (kWh)", "化成分容耗电 (kWh)",
                       "模组焊接耗电 (kWh)", "Pack装配耗电 (kWh)"],
        "运输": ["厂内周转运输 (tkm)"]
    },
    "3.5.3 分销和储存阶段": {
        "物料与能耗": ["仓储温控耗电 (kWh)"],
        "运输": ["电池仓储周转运输 (tkm)", "汽车动力电池成品运输 (tkm)"]
    },
    "3.5.4 产品使用阶段": {
        "物料与能耗": ["车辆行驶充电耗电 (kWh)", "电池热管理系统耗电 (kWh)"],
        "运输": ["售后维保运输 (tkm)"]
    },
    "3.5.5 废弃处置阶段": {
        "物料与能耗": ["动力电池回收拆解耗电 (kWh)", "回收清洗废水 (t)", "废弃正极材料 (kg)", "废弃负极材料 (kg)",
                       "废弃电解液 (kg)", "废弃隔膜 (kg)", "废弃铝箔 (kg)", "废弃铜箔 (kg)", "废弃电池结构胶 (kg)",
                       "废弃绝缘材料 (kg)", "废弃BMS (kg)", "废弃铝合金壳体 (kg)", "废弃水冷板 (kg)", "废弃铜排 (kg)"],
        "运输": ["废弃电池运输 (tkm)", "废弃包装材料运输 (tkm)"]
    }
}

user_records = {}
results = {stage: 0.0 for stage in UI_STRUCTURE.keys()}

# 渲染预设项目的填报
for stage_name, categories in UI_STRUCTURE.items():
    user_records[stage_name] = {"Material": [], "Transport": []}
    with st.expander(f"⚙️ 展开录入预设清单：{stage_name}", expanded=False):
        for cat_name, items in categories.items():
            st.markdown(f"**📍 {cat_name}**")
            cols = st.columns(3)
            for i, item in enumerate(items):
                with cols[i % 3]:
                    val = st.number_input(item, min_value=0.0, step=1.0, key=f"{stage_name}_{item}")
                    match = re.search(r"(.+)\s*\((.+)\)", item)
                    name = match.group(1).strip() if match else item
                    unit = match.group(2).strip() if match else "-"
                    if "运输" in cat_name or "运输" in name:
                        factor = FACTOR_DB.get(item, default_trans_factor)
                        user_records[stage_name]["Transport"].append((name, val, unit, "运输"))
                    else:
                        factor = FACTOR_DB.get(item, 0.0)
                        user_records[stage_name]["Material"].append((name, val, unit, "主料/辅料/能耗"))
                    results[stage_name] += val * factor
            st.divider()

# ==========================================
# 🌟 核心新功能：补充自定义排放清单 (交互式表格)
# ==========================================
st.markdown("### ➕ 补充自定义排放清单")
st.info(
    "如果实际生产中存在上述预设清单未涵盖的物料、能耗或运输过程，请在下方表格点击空白行手动添加。系统会自动计入对应阶段。")

stage_options = list(UI_STRUCTURE.keys())
custom_df_init = pd.DataFrame(
    columns=["所属生命周期阶段", "排放源名称", "消耗数量", "单位", "排放因子数值 (kgCO2e/单位)"])

# 渲染可交互的数据表格
custom_inputs = st.data_editor(
    custom_df_init,
    column_config={
        "所属生命周期阶段": st.column_config.SelectboxColumn("所属生命周期阶段", options=stage_options, required=True),
        "排放源名称": st.column_config.TextColumn("排放源名称 (如: 新型溶剂)", required=True),
        "消耗数量": st.column_config.NumberColumn("消耗数量", min_value=0.0, default=0.0, required=True),
        "单位": st.column_config.TextColumn("单位 (如: kg, L)", required=True),
        "排放因子数值 (kgCO2e/单位)": st.column_config.NumberColumn("排放因子数值", default=0.0, required=True),
    },
    num_rows="dynamic",
    use_container_width=True,
    key="custom_data_editor"
)

# 处理用户手填的自定义数据
for index, row in custom_inputs.iterrows():
    c_stage = row.get("所属生命周期阶段")
    c_name = row.get("排放源名称")
    c_val = row.get("消耗数量")
    c_unit = row.get("单位")
    c_factor = row.get("排放因子数值 (kgCO2e/单位)")

    # 确保填完整了才计算，防止报错
    if pd.notna(c_stage) and pd.notna(c_name) and pd.notna(c_val) and pd.notna(c_factor):
        if c_stage in results:
            results[c_stage] += c_val * c_factor
            # 标记类型为“自定义补充”，它会自动被后面的 Word 生成器抓取并写入表格
            user_records[c_stage]["Material"].append((c_name, c_val, c_unit if pd.notna(c_unit) else "-", "自定义补充"))


# Word 表格辅助函数
def add_word_table(doc, title, headers, data_rows):
    doc.add_paragraph(title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r in data_rows:
        row = table.add_row().cells
        for i, val in enumerate(r):
            row[i].text = str(val)


# ==========================================
# 🌟 第三步：计算与极致排版的 Word 导出
# ==========================================
st.markdown("<br>", unsafe_allow_html=True)

if st.button("🚀 提交全部数据，生成极致排版规范报告", type="primary", use_container_width=True):
    if not company_intro.strip() or not product_intro.strip() or not production_process.strip():
        st.error("❌ 警告：您有必填项未完成！请先在上方【Part 1: 概述】中填写企业简介、产品介绍和产品生产工艺。")
    else:
        total_carbon = sum(results.values())
        st.success("✅ 数据核算完成！")

        st.markdown("### 📊 本次测算碳足迹概览")
        st.metric(label="🌟 生命周期总碳足迹 (kgCO2e)", value=f"{total_carbon:,.6f}")

        col_chart1, col_chart2 = st.columns([1.2, 2])
        res_data = []
        for stage, em in results.items():
            pct = f"{(em / total_carbon * 100):.4f}" if total_carbon > 0 else "0.0000"
            res_data.append([stage.split(" ")[1], em, f"{pct}%"])

        df_results = pd.DataFrame(res_data, columns=["生命周期阶段", "排放量(kgCO2e)", "占比(%)"])
        with col_chart1:
            st.write("**各阶段排放明细表**")
            df_display = df_results.copy()
            df_display["排放量(kgCO2e)"] = df_display["排放量(kgCO2e)"].apply(lambda x: f"{x:,.6f}")
            st.dataframe(df_display, hide_index=True)
        with col_chart2:
            st.write("**各阶段碳排放量动态柱状图**")
            st.bar_chart(df_results.set_index("生命周期阶段")["排放量(kgCO2e)"])

        st.info("报告文件排版已完成，请点击下方按钮下载。")

        # ==================================
        # 开始构建 Word 文档 (引入全套强管控排版)
        # ==================================
        doc = Document()

        # 🎯 全局强制字体与样式设定
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Times New Roman'
        style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style_normal.font.size = Pt(12)
        style_normal.font.color.rgb = RGBColor(0, 0, 0)

        for i in range(1, 4):
            h_style = doc.styles[f'Heading {i}']
            h_style.font.name = 'Times New Roman'
            h_style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            h_style.font.color.rgb = RGBColor(0, 0, 0)
            if i == 1:
                h_style.font.size = Pt(16)
                h_style.font.bold = True
            elif i == 2:
                h_style.font.size = Pt(14)
                h_style.font.bold = True
            elif i == 3:
                h_style.font.size = Pt(12)
                h_style.font.bold = True

        # --- 封面生成 ---
        for _ in range(5): doc.add_paragraph()
        title = doc.add_paragraph()
        title_run = title.add_run('产品碳足迹评价报告')
        title_run.font.size = Pt(26)
        title_run.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        sub_run = subtitle.add_run('——汽车动力电池，100kwh\n\n\n')
        sub_run.font.size = Pt(16)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info = doc.add_paragraph()
        info_run = info.add_run(f'报告单位： 体验账号\n编制日期： {datetime.date.today().strftime("%Y年%m月%d日")}')
        info_run.font.size = Pt(14)
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

        # --- 🎯 极致还原的目录页生成 ---
        doc.add_heading('目录', level=1)

        toc_lines = [
            ("1. 概述", "1", 0),
            ("1.1 企业简介", "1", 1),
            ("1.2 产品介绍", "1", 1),
            ("1.3 产品生产工艺", "1", 1),
            ("1.4 评价依据和要求", "1", 1),
            ("2. 目的和范围定义", "2", 0),
            ("2.1 研究目的", "2", 1),
            ("2.2 功能单位", "2", 1),
            ("2.3 时间范围", "2", 1),
            ("2.4 系统边界", "2", 1),
            ("2.5 温室气体种类", "2", 1),
            ("2.6 取舍原则", "3", 1),
            ("3. 生命周期数据清单分析", "4", 0),
            ("3.1 数据质量要求", "4", 1),
            ("3.2 碳足迹计算方法", "5", 1),
            ("3.3 分配", "6", 1),
            ("3.4 假设", "7", 1),
            ("3.5 生命周期清单数据", "7", 1),
            ("3.5.1 原材料获取阶段", "7", 2),
            ("3.5.2 生产制造阶段", "8", 2),
            ("3.5.3 分销和储存阶段", "8", 2),
            ("3.5.4 产品使用阶段", "8", 2),
            ("3.5.5 废弃处置阶段", "9", 2),
            ("4. 产品碳足迹评价结果", "10", 0),
            ("5. 不确定性分析", "11", 0),
            ("5.1 不确定性分析方法", "11", 1),
            ("5.2 不确定性分析结果", "12", 1),
            ("6. 结论", "14", 0),
            ("附录：排放因子选择表", "15", 0),
            ("参考文献", "16", 0)
        ]

        for text, page, level in toc_lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(level * 20)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

            run_title = p.add_run(text)
            if level == 0: run_title.bold = True
            p.add_run('\t')
            run_page = p.add_run(page)
            if level == 0: run_page.bold = True

        doc.add_page_break()

        # --- 正文输出 ---
        doc.add_heading('1. 概述', level=1)
        doc.add_heading('1.1 企业简介', level=2)
        doc.add_paragraph(company_intro)
        doc.add_heading('1.2 产品介绍', level=2)
        doc.add_paragraph(product_intro)
        doc.add_heading('1.3 产品生产工艺', level=2)
        doc.add_paragraph(production_process)
        doc.add_heading('1.4 评价依据和要求', level=2)
        doc.add_paragraph(evaluation_basis if evaluation_basis.strip() else def_p1_4)

        doc.add_heading('2. 目的和范围定义', level=1)
        doc.add_heading('2.1 研究目的', level=2)
        doc.add_paragraph(purpose if purpose.strip() else def_p2_1)
        doc.add_heading('2.2 功能单位', level=2)
        doc.add_paragraph(functional_unit if functional_unit.strip() else def_p2_2)
        doc.add_heading('2.3 时间范围', level=2)
        doc.add_paragraph(time_scope if time_scope.strip() else def_p2_3)
        doc.add_heading('2.4 系统边界', level=2)
        doc.add_paragraph(system_boundary if system_boundary.strip() else def_p2_4)
        doc.add_heading('2.5 温室气体种类', level=2)
        doc.add_paragraph(ghg_types if ghg_types.strip() else def_p2_5)
        doc.add_heading('2.6 取舍原则', level=2)
        doc.add_paragraph(cutoff_rules if cutoff_rules.strip() else def_p2_6)

        doc.add_heading('3. 生命周期数据清单分析', level=1)
        doc.add_heading('3.1 数据质量要求', level=2)
        doc.add_paragraph(
            '综合ISO 14067：2018、PAS 2050：2011对数据质量的要求，本次评价产品碳足迹的活动水平和排放因子满足相关性、完整性、一致性、连贯性、透明性、准确性及避免重复计算的要求。')
        doc.add_heading('3.2 碳足迹计算方法', level=2)
        doc.add_paragraph('采用公式：E = Σ (AD × EF × GWP) 进行计算。')
        doc.add_heading('3.3 分配', level=2)
        doc.add_paragraph(
            '在无法避免分配的情况下，系统的输入和输出应以反应它们之间潜在的物理关系的方式，在其不同的产品或功能之间进行划分。')
        doc.add_heading('3.4 假设', level=2)
        doc.add_paragraph('基于数据可得性与碳足迹核算需求的综合评估，已在合理假设下完成测算。')
        doc.add_heading('3.5 生命周期清单数据', level=2)

        table_idx = 2
        for stage_name, records in user_records.items():
            pure_stage_name = stage_name.split(' ')[1]
            doc.add_heading(stage_name, level=3)
            if records["Material"]:
                doc.add_paragraph(f'{pure_stage_name}的清单数据，如表3-{table_idx}所示。')
                # 预设数据和手填补充数据都会在这里打印出来！
                add_word_table(doc, f'表 3-{table_idx}：{pure_stage_name}数据清单', ['名称', '消耗数量', '单位', '类型'],
                               records["Material"])
                table_idx += 1
            if records["Transport"]:
                trans_formatted = [[row[0], "1", "道路运输", "-", row[1], "-"] for row in records["Transport"]]
                doc.add_paragraph(f'{pure_stage_name}的运输清单数据，如表3-{table_idx}所示。')
                add_word_table(doc, f'表 3-{table_idx}：{pure_stage_name}运输数据清单',
                               ['运输物', '路段', '运输方式', '货物重量', '运输里程', '能源消耗'], trans_formatted)
                table_idx += 1

        doc.add_heading('4. 产品碳足迹评价结果', level=1)
        doc.add_paragraph(
            f'通过计算，功能单位（1个汽车动力电池）的全生命周期碳足迹为 {total_carbon:,.6f} kgCO2e，单位产品排放量为 {total_carbon:,.6f} kgCO2e/个，碳足迹的整体情况，如表4-1所示。')
        res_data_word = [[r[0], f"{r[1]:,.6f}", r[2]] for r in res_data]
        res_data_word.append(['总计', f"{total_carbon:,.6f}", '100%'])
        add_word_table(doc, '表4-1：产品碳足迹评价结果', ['生命周期阶段', '排放量（kgCO2e）', '占比（%）'], res_data_word)

        doc.add_heading('5. 不确定性分析', level=1)
        doc.add_heading('5.1 不确定性分析方法', level=2)
        doc.add_paragraph(
            '产品碳足迹评价的不确定性分析，采用定性分析法，包括活动水平数据质量评级、排放因子的质量评级。活动水平数据质量评级，如表5-1所示。')
        add_word_table(doc, '表5-1：活动水平数据质量评级', ['质量等级', '描述', '分值'],
                       [['好', '量测值：实际量测数值...', '5'], ['较好', '计算值：以某合理方法进行计算的数值...', '3'],
                        ['一般', '理论值/经验值：根据理论推导...', '2'], ['差', '参考文献：由其它文献取得...', '1']])
        doc.add_paragraph(
            '排放因子质量评级，从时间相关性、地域相关性、技术相关性、数据准确度、方法学等方面评定，具体标准，如表5-2、5-3、5-4、5-5、5-6所示。')
        add_word_table(doc, '表5-2：排放因子的质量评级-时间相关性', ['时间相关性', '分值'],
                       [['<5年', '5'], ['5–10年', '3'], ['10–15年', '2'], ['>15年（及未知年份）', '1']])
        add_word_table(doc, '表5-3：排放因子的质量评级-地域相关性', ['地域相关性', '分值'],
                       [['完全符合所评估产品生产地点', '5'], ['数据为国家层面的数据', '3'],
                        ['数据为全球平均数据', '1']])
        add_word_table(doc, '表5-4：排放因子的质量评级-技术相关性', ['技术相关性', '分值'],
                       [['完全符合所评估产品生产技术', '5'], ['行业平均数据', '3'], ['替代数据', '1']])
        add_word_table(doc, '表5-5：排放因子的质量评级-数据准确度', ['数据准确度', '分值'],
                       [['变异性低', '5'], ['变异性未量化，考虑为较低', '3'], ['变异性未量化，考虑为较高', '2'],
                        ['变异性高', '1']])
        add_word_table(doc, '表5-6：排放因子的质量评级-方法学', ['方法学的适合性及一致性', '分值'],
                       [['PAS 2050/补充要求所规定的排放因子', '5'], ['政府/国际政府组织/行业发布的排放因子', '3'],
                        ['公司/其他机构发布的排放因子', '1']])
        doc.add_heading('5.2 不确定性分析结果', level=2)
        doc.add_paragraph('汽车动力电池产品碳足迹评价的活动数据和排放因子数据质量分析结果，如表5-7、5-8所示。')
        add_word_table(doc, '表5-7：活动数据质量分析结果', ['活动数据类别', '活动数据描述', '质量级别', '得分'],
                       [['原材料类', '', '', ''], ['能源资源类', '', '', ''], ['运输类', '', '', ''],
                        ['产品使用类', '', '', ''], ['废弃处置类', '', '', '']])
        add_word_table(doc, '表5-8：排放因子数据质量分析结果', ['排放因子类别', '排放因子描述', '平均得分'],
                       [['原材料类', '', ''], ['能源资源类', '', ''], ['运输类', '', ''], ['产品使用类', '', ''],
                        ['废弃处置类', '', '']])

        doc.add_heading('6. 结论', level=1)
        doc.add_paragraph(
            f'体验账号汽车动力电池的产品碳足迹评价，涵盖的时间范围是2026年01月01日至2026年12月31日。功能单位（1个汽车动力电池）的全生命周期碳足迹为 {total_carbon:,.6f} kgCO2e，其中原材料获取、生产制造、分销储存、产品使用、废弃处置阶段的排放占比分别为：{res_data[0][2]}、{res_data[1][2]}、{res_data[2][2]}、{res_data[3][2]}、{res_data[4][2]}。')

        doc.add_page_break()
        doc.add_heading('附录：排放因子选择表', level=1)
        full_appendix_data = [
            ['废弃电池结构胶', '电池 - 镍氢电池', '28380', 'kgCO₂e/t', '英国环境、食品及农村事务部 (DEFRA)'],
            ['废弃绝缘材料', '铸铁材料', '1.82', 'kgCO₂e/kg', '中国产品全生命周期温室气体排放系数库'],
            ['回收清洗废水', '废水处理', '0.8581', 'kgCO₂e/m³', '澳大利亚生命周期清单数据库'],
            ['动力电池回收拆解耗电', '电池 - 镍氢电池', '28380', 'kgCO₂e/t', '英国环境、食品及农村事务部 (DEFRA)'],
            ['车辆行驶充电耗电', '汽车', '207.8', 'gCO₂e/km', '中国产品全生命周期温室气体排放系数库'],
            ['电池热管理系统耗电', '电池 - 镍氢电池', '28380', 'kgCO₂e/t', '英国环境、食品及农村事务部 (DEFRA)'],
            ['仓储温控耗电', '仓储', '24.404', 'kgCO₂e/㎡', '碳阻迹（北京）科技有限公司'],
            ['有机废气', '煤矿废气', '56.8', 'kgCO₂e/GJ', '澳大利亚气候变化和能源部'],
            ['光伏发电电力', '电力碳足迹', '0.65', 'kgCO₂e/度', '台湾“行政院环境保护署”'],
            ['电芯涂布烘干耗电', '涂布纸', '0.7999', 'kgCO₂e/kg', '澳大利亚生命周期清单数据库'],
            ['辊压工序耗电', '炼焦工序', '520', 'kgCO₂e/t', '哈尔滨工业大学'],
            ['电芯注液耗电', '液氧', '0.39', 'kgCO₂e/kg', '台湾“行政院环境保护署”'],
            ['化成分容耗电', '商业用电', '29.08', 'tCO₂e/百万日元', '日本环境省'],
            ['模组焊接耗电', '焊接和焊接设备制造', '0.249', 'kgCO₂e/美元', '美国环保署'],
            ['Pack装配耗电', '商业用电', '468', 'tCO₂e/百万kWh', '日本环境省'],
            ['纸质护角', '纸质包装面包片', '1.131', 'kgCO₂e/个', '中国产品全生命周期温室气体排放系数库'],
            ['电池结构胶', '电池 - 镍氢电池', '28380', 'kgCO₂e/t', '英国环境、食品及农村事务部 (DEFRA)'],
            ['绝缘材料', '隔热（或隔音、绝缘）材料', '1852.0809', 'kgCO₂e/t', '英国环境、食品及农村事务部 (DEFRA)'],
            ['电解质锂盐', '盐-零售', '0.01', 'kgCO₂e/kg', '丹麦绿色智库CONCITO'],
            ['导电剂', '灭菌剂', '3.9', 'tCO₂e/t', '四川大学'],
            ['粘结剂', '明胶，粘合剂', '0.00223', 'tCO₂e/kg', '日本环境省'],
            ['主料2', '磷酸铁锂', '25', 'kgCO₂e/kg', '中国产品全生命周期温室气体排放系数库'],
            ['主料1', '电池', '6308', 'kgCO₂e/t', '英国环境、食品及农村事务部 (DEFRA)'],
            ['废弃正极材料', '铸铁材料', '1.82', 'kgCO₂e/kg', 'China Products Carbon Footprint Factors Database'],
            ['废弃负极材料', '铸铁材料', '1.82', 'kgCO₂e/kg', 'China Products Carbon Footprint Factors Database'],
            ['废弃电解液', '电解液：六氟磷酸锂', '19.6', 'kgCO₂e/kg', 'China Automotive Data Co. Ltd'],
            ['废弃隔膜', '再生料-废容器-PET废塑胶片', '0.39', 'kgCO₂e/kg', 'Carbon Footprint Taiwan MOEVN 20'],
            ['废弃铝箔', '铝箔', '2.39', 'kgCO₂e/kg', 'Korea Environmental Industry Technology Research Institute'],
            ['废弃铜箔', '铜箔', '12.4', 'kgCO₂e/kg', 'Taiwan Environmental Protection Agency'],
            ['废弃电池管理系统（BMS）', '电池 - 镍氢电池', '28380', 'kgCO₂e/t', 'Department for Environment'],
            ['废弃电池壳体（铝合金）', '电池 - 镍氢电池', '28380', 'kgCO₂e/t', 'Department for Environment'],
            ['废弃冷却系统（水冷板）', '自来水', '0.167', 'kgCO₂e/m³', 'Taiwan Environmental Protection Agency'],
            ['废弃连接件（铜排）', '再生料-废铜', '0.14', 'kgCO₂e/kg', 'Carbon Footprint Taiwan MOEVN 20'],
            ['电力', '电力', '0.6205', 'kgCO₂e/kWh', 'Ministry of Ecology and Environment'],
            ['废水', '工业废水', '0.1185', 'kgCO₂e/t', 'New Zealand Ministry of the Environment'],
            ['水', '水', '0.344', 'kgCO₂e/m³', 'Department for Environment'],
            ['天然气', '天然气', '2.06672', 'kgCO₂e/m³', 'Department for Environment'],
            ['木质托盘', '木质托盘48x40', '1.28', 'kgCO₂e/kg', 'Scientific Data'],
            ['塑料薄膜', '塑料：普通塑料', '3172.49932', 'kgCO₂e/t', 'Department for Environment'],
            ['电池管理系统（BMS）', '电池 - 镍氢电池', '28380', 'kgCO₂e/t', 'Department for Environment'],
            ['电池壳体（铝合金）', '电池 - 镍氢电池', '28380', 'kgCO₂e/t', 'Department for Environment'],
            ['冷却系统（水冷板）', '自来水', '0.167', 'kgCO₂e/m³', 'Taiwan Environmental Protection Agency'],
            ['连接件（铜排）', '铜箔', '12.4', 'kgCO₂e/kg', 'Carbon Footprint Taiwan MOEVN 20'],
            ['正极材料（磷酸铁锂）', '铸铁材料', '1.82', 'kgCO₂e/kg', 'China Products Carbon Footprint Factors Database'],
            ['负极材料（石墨）', '石墨', '5.5', 'kgCO₂e/kg', 'China Automotive Data Co. Ltd'],
            ['电解液', '电解液：六氟磷酸锂', '19.6', 'kgCO₂e/kg', 'China Automotive Data Co. Ltd'],
            ['隔膜', '塑料薄膜包装袋', '3.24', 'kgCO₂e/kg', 'China Products Carbon Footprint Factors Database'],
            ['铝箔', '铝箔', '2.39', 'kgCO₂e/kg', 'Korea Environmental Industry Technology Research Institute'],
            ['铜箔', '铜箔', '12.4', 'kgCO₂e/kg', 'Taiwan Environmental Protection Agency'],
            ['电池仓储周转运输', '重型柴油货车运输（载重30t）', '0.078', 'kgCO₂e/吨公里',
             '中华人民共和国住房和城乡建设部'],
            ['正极材料运输', '货运火车', '0.02782', 'kgCO₂e/tkm', '英国环境、食品及农村事务部 (DEFRA)'],
            ['负极材料运输', '货运火车', '0.02782', 'kgCO₂e/tkm', '英国环境、食品及农村事务部 (DEFRA)'],
            ['电解液运输', '重型柴油货车运输（载重30t）', '0.078', 'kgCO₂e/吨公里', '中华人民共和国住房和城乡建设部'],
            ['隔膜运输', '重型柴油货车运输（载重30t）', '0.078', 'kgCO₂e/吨公里', '中华人民共和国住房和城乡建设部'],
            ['铝箔运输', '中型柴油货车运输（载重8t）', '0.179', 'kgCO₂e/吨公里', '中华人民共和国住房和城乡建设部'],
            ['铜箔运输', '重型柴油货车运输（载重30t）', '0.078', 'kgCO₂e/吨公里', '中华人民共和国住房和城乡建设部'],
            ['辅料运输', '重型柴油货车运输（载重30t）', '0.078', 'kgCO₂e/吨公里', '中华人民共和国住房和城乡建设部'],
            ['废弃电池运输', '重型柴油货车运输（载重30t）', '0.078', 'kgCO₂e/tkm', '中华人民共和国住房和城乡建设部'],
            ['废弃包装材料运输', '重型柴油货车运输（载重30t）', '0.078', 'kgCO₂e/tkm', '中华人民共和国住房和城乡建设部'],
            ['汽车动力电池成品运输', '重型柴油货车运输（载重30t）', '0.078', 'kgCO₂e/tkm',
             '中华人民共和国住房和城乡建设部']
        ]
        add_word_table(doc, '附表：系统引用的主要排放因子',
                       ['排放源名称', '排放因子名称', '排放因子数值', '排放因子单位', '排放因子来源'],
                       full_appendix_data)
        doc.add_paragraph(
            '\n参考文献：\n[1] 《商品和服务在生命周期内的温室气体排放评价规范》（PAS 2050:2011）\n[2] 《温室气体—产品碳足迹—量化要求与指南》（ISO 14067:2018）\n[3] 《温室气体核算体系：产品寿命周期核算与报告标准》\n[4] 《IPCC WGI Sixth Assessment Report》')

        bio = io.BytesIO()
        doc.save(bio)

        st.download_button(
            label="📥 点击下载极致排版 Word 测算报告",
            data=bio.getvalue(),
            file_name="产品碳足迹评价报告_极致排版版.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )